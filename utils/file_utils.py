import io
import docx
import fitz
import json
import base64
import chardet
import pdfplumber
import pandas as pd
from PIL import Image
from flask import current_app
import tiktoken


class FileOperation:

    @staticmethod
    def extract_images_from_pdf(pdf_path, is_content):
        if not is_content:
            return []
        with fitz.open(stream=pdf_path) as doc:
            base64_img = []
            for page_index in range(len(doc)):
                if page_index not in is_content:
                    continue
                # if img:= doc[page_index].get_images(full=True):
                #     xref = img[-1][0]
                #     base_image = doc.extract_image(xref)
                #     image_bytes = base_image["image"]
                #     img = Image.open(io.BytesIO(image_bytes))
                #     buffered = io.BytesIO()
                #     img.save(buffered, format="PNG")
                #     img = base64.b64encode(buffered.getvalue()).decode()
                #     base64_img.append(img)
                # else:
                # page = doc[page_index]
                # pix = page.get_pixmap(dpi=500)
                # img = Image.open(io.BytesIO(pix.tobytes("png")))
                # buffered = io.BytesIO()
                # img.save(buffered, format="PNG")
                # img = base64.b64encode(buffered.getvalue()).decode()
                # base64_img.append(img)
                # 建议改 dpi 为 150 或用缩放矩阵控制分辨率
                page = doc[page_index]
                pix = page.get_pixmap(dpi=150)

                # 直接拿 PNG 字节，不要再二次转换
                img_bytes = pix.tobytes("png")

                # 转 base64
                img_b64 = base64.b64encode(img_bytes).decode()
                base64_img.append(img_b64)
        return base64_img

    @staticmethod
    def extract_text_from_pdf(pdf_path):
        pdf_bytes = io.BytesIO(pdf_path)
        is_content = []
        final_text = []
        tables = []
        df = ""
        with pdfplumber.open(pdf_bytes) as pdf:
            for page in pdf.pages:
                table = page.extract_table()
                if table and any(any(once) for once in table):
                    tables.append(table)
                    table_bboxes = [pos.bbox for pos in page.find_tables()]
                    filtered_text = page.extract_words()
                    for word in filtered_text:
                        x0, y0, x1, y1 = word["x0"], word["top"], word["x1"], word["bottom"]
                        inside_table = any(
                            t_x0 <= x0 <= t_x1 and t_y0 <= y0 <= t_y1
                            for (t_x0, t_y0, t_x1, t_y1) in table_bboxes
                        )
                        if not inside_table:
                            final_text.append(word["text"])
                else:
                    text = page.extract_text()
                    if text and len(text) > 10:
                        final_text.append(text)
                    else:
                        is_content.append(page.page_number - 1)
            if tables:
                try:
                    df = "\n".join(pd.DataFrame(table[1:], columns=table[0]).to_json(force_ascii=False) for table in tables)
                except:
                    df = json.dumps(tables)
            if final_text:
                final_text = " ".join(final_text) + "\n"
            else:
                final_text = ""
        return final_text + df, is_content

    @staticmethod
    def extract_text_from_word(docx_path):
        doc = docx.Document(docx_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        tables = []
        df = ""

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)
        if tables:
            try:
                df = pd.DataFrame(tables[0][1:], columns=tables[0][0]).to_json(force_ascii=False)
            except:
                df = json.dumps(tables, ensure_ascii=False)
        return text.strip() + df

    @staticmethod
    def extract_images_from_word(docx_path):
        doc = docx.Document(docx_path)
        base64_img = []
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                image = doc.part.rels[rel].target_part.blob
                img = Image.open(io.BytesIO(image))
                buffered = io.BytesIO()
                img.save(buffered, format="PNG")
                # save pages
                img_b64 = base64.b64encode(buffered.getvalue()).decode()
                base64_img.append(img_b64)
        return base64_img

    @staticmethod
    def check_pdf(pdf_path):
        with fitz.open(stream=pdf_path) as doc:
            for page_index in range(len(doc)):
                if doc[page_index].get_images(full=True) or not doc[page_index].get_text("text"):
                    page = doc[page_index]
                    pix = page.get_pixmap(dpi=500)
                    rect = page.rect
                    page.clean_contents()
                    page.insert_image(rect, pixmap=pix)
            pdf_buffer = io.BytesIO()
            doc.save(pdf_buffer, garbage=4, deflate=True)
            pdf_byte_path = pdf_buffer.getvalue()
        return pdf_byte_path

    @staticmethod
    def extract_picture(picture_path):
        base64_img = []
        img = Image.open(picture_path)
        buffered = io.BytesIO()
        img.save(buffered, format="PNG")
        img = base64.b64encode(buffered.getvalue()).decode()
        base64_img.append(img)
        return base64_img

    def __call__(self, username: str, attachment_names: list):
        if not isinstance(attachment_names, list):
            return {"message": "Invalid attachment_names format", "status": 400}
        else:
            results = {}
            for attachment_name in attachment_names:
                file_extension = attachment_name.rsplit(".", 1)[1].lower()
                blob_client = current_app.container_client.get_blob_client(f"{username}/{attachment_name}")
                stream = blob_client.download_blob().readall()
                file_stream = io.BytesIO(stream)
                encoding = chardet.detect(stream)["encoding"]

                if file_extension == "txt":
                    results[attachment_name] = {
                        "text": stream.decode(encoding),
                        "images": []
                    }

                elif file_extension == "csv":
                    df = pd.read_csv(file_stream, encoding=encoding)
                    results[attachment_name] = {
                        "text": df.to_json(force_ascii=False),
                        "images": []
                    }

                elif file_extension == "json":
                    results[attachment_name] = {
                        "text": stream.decode(encoding),
                        "images": []
                    }

                elif file_extension in ["xlsx", "xls"]:
                    try:
                        if file_extension == "xlsx":
                            df = pd.read_excel(file_stream, engine="openpyxl")
                        elif file_extension == "xls":
                            # 需要安装 pip install xlrd==1.2.0
                            df = pd.read_excel(file_stream, engine="xlrd")
                        else:
                            raise ValueError("Unsupported Excel file type")
                        results[attachment_name] = {
                            "text": df.to_json(force_ascii=False),
                            "images": []
                        }
                    except Exception as e:
                        results[attachment_name] = {
                            "text": f"Failed to read Excel file: {str(e)}",
                            "images": []
                        }

                # elif file_extension == "pdf":
                #     pdf_text, page_num = self.extract_text_from_pdf(stream)
                #     pdf_images = self.extract_images_from_pdf(stream, page_num)
                #     results[attachment_name] = {
                #         "text": pdf_text,
                #         "images": pdf_images
                #     }
                elif file_extension == "pdf":
                    pdf_stream = io.BytesIO(stream)  # 二进制转 BytesIO
                    pdf_text, page_num = self.extract_text_from_pdf(stream)
                    pdf_images = self.extract_images_from_pdf(pdf_stream, page_num)
                    results[attachment_name] = {
                        "text": pdf_text,
                        "images": pdf_images
                    }

                elif file_extension == "docx":
                    word_text = self.extract_text_from_word(file_stream)
                    try:
                        word_images = self.extract_images_from_word(file_stream)
                    except:
                        word_images = []
                    results[attachment_name] = {
                        "text": word_text,
                        "images": word_images
                    }

                elif file_extension in ["jpg", "jpeg", "png"]:
                    images = self.extract_picture(file_stream)
                    results[attachment_name] = {
                        "text": "",
                        "images": images
                    }

                else:
                    results[attachment_name] = {
                        "text": f"Unsupported file type: {file_extension}",
                        "images": []
                    }

        return results


# ========================
# 全局缓存与常量设置
# ========================
_token_cache = {}
_cache_max_size = 5000000

MODEL_TOKEN_LIMIT = {
    "gpt-4o": 600000,
    "gpt-4o-mini": 600000,
    "gpt-4-turbo": 600000,
    "gpt-35-turbo": 16384,
    "gpt-3.5-turbo": 16384,
}

# ========================
# 主入口函数
# ========================
def cal_tokens(username: str, attachment_names: list, deploy_model: str = "gpt-4o"):
    """
    快速计算文件的token数量 - 优化版本
    Args:
        username: 用户名
        attachment_names: 文件名列表
        deploy_model: 模型名称，默认 gpt-4o
    Returns:
        dict: {"total_tokens": int, "file_tokens": {filename: int, ...}, "limit": int}
    """
    if not isinstance(attachment_names, list):
        return {"error": "Invalid attachment_names format", "total_tokens": 0}

    try:
        try:
            encoding = tiktoken.encoding_for_model(deploy_model)
        except KeyError:
            encoding = tiktoken.encoding_for_model("gpt-4")

        total_tokens = 0
        file_tokens = {}

        model_limit = MODEL_TOKEN_LIMIT.get(deploy_model, 600000)

        for attachment_name in attachment_names:
            try:
                cache_key = f"{username}:{attachment_name}:{deploy_model}"
                if cache_key in _token_cache:
                    tokens = _token_cache[cache_key]
                    file_tokens[attachment_name] = tokens
                    total_tokens += tokens
                    continue

                file_extension = attachment_name.rsplit(".", 1)[1].lower()
                blob_client = current_app.container_client.get_blob_client(f"{username}/{attachment_name}")

                tokens = _estimate_tokens_fast(blob_client, file_extension, encoding)

                # 缓存结果
                _cache_with_limit(cache_key, tokens)
                file_tokens[attachment_name] = tokens
                total_tokens += tokens

            except Exception as e:
                file_tokens[attachment_name] = 0
                print(f"⚠️ Error processing {attachment_name}: {e}")

        return {
            "total_tokens": total_tokens,
            "file_tokens": file_tokens,
            "limit": model_limit,
            "within_limit": total_tokens <= model_limit,
            "success": True,
        }

    except Exception as e:
        return {"error": str(e), "total_tokens": 0, "success": False}

# ========================
# 缓存函数
# ========================
def _cache_with_limit(key: str, value: int):
    """带上限的LRU缓存"""
    global _token_cache
    if len(_token_cache) >= _cache_max_size:
        keys_to_remove = list(_token_cache.keys())[: _cache_max_size // 4]  # 仅删1/4
        for k in keys_to_remove:
            _token_cache.pop(k, None)
    _token_cache[key] = value

# ========================
# 改进版快速估算函数
# ========================
def _estimate_tokens_fast(blob_client, file_extension: str, encoding):
    """
    改进版：更精确估算 PDF（含 base64 图像流）的 token 数量
    """
    try:
        blob_properties = blob_client.get_blob_properties()
        file_size = blob_properties.size  # bytes

        # === 图片类文件 ===
        if file_extension in ["jpg", "jpeg", "png"]:
            if file_size < 100 * 1024:
                return 100
            elif file_size < 500 * 1024:
                return 200
            else:
                return 300

        # === PDF 文件（重点优化） ===
        if file_extension == "pdf":
            # 下载前 512KB 样本进行结构分析
            sample_len = min(512 * 1024, file_size)
            sample_data = blob_client.download_blob(offset=0, length=sample_len).readall()

            # 尝试读取 PDF 结构（部分加载即可）
            try:
                doc = fitz.open(stream=sample_data, filetype="pdf")
                n_pages = len(doc)
                n_images = 0
                for page in doc:
                    n_images += len(page.get_images(full=True))
                doc.close()
            except Exception:
                n_pages, n_images = 1, 0  # 解析失败，保守估算

            # 图像页比例
            img_ratio = min(n_images / max(n_pages, 1), 1.0)

            # === 估算逻辑 ===
            # 1 token ≈ 4 字节（纯文本页）
            # 图片页约等价于 base64 转换后每 3 字节→4字节，约 ×1.33 token 消耗
            base_text_ratio = 1 - img_ratio
            base_image_ratio = img_ratio * 1.33

            # 平均有效比例（非线性缓冲）
            effective_ratio = 4 * base_text_ratio + 6 * base_image_ratio  # 偏大以防低估

            # 经验衰减系数：越大文件→单位token/byte越低（考虑压缩率）
            decay = 1.0
            if file_size > 1 * 1024 * 1024:  # >1MB
                decay = 0.6
            if file_size > 3 * 1024 * 1024:  # >3MB
                decay = 0.5

            tokens = int(file_size / effective_ratio * decay)

            # 限制范围（最小2k，最大120k）
            tokens = int(file_size / effective_ratio * decay)
            return tokens

        # === 文本类（TXT/JSON/CSV） ===
        if file_extension in ["txt", "json", "csv"]:
            sample_size = min(4096, file_size)
            sample_data = blob_client.download_blob(offset=0, length=sample_size).readall()
            encoding_type = chardet.detect(sample_data)["encoding"] or "utf-8"
            sample_text = sample_data.decode(encoding_type, errors="ignore")
            if not sample_text:
                return int(file_size / 4)
            sample_tokens = len(encoding.encode(sample_text))
            token_density = sample_tokens / max(len(sample_text), 1)
            avg_bytes_per_char = max(sample_size / max(len(sample_text), 1), 1.0)
            estimated_chars = file_size / avg_bytes_per_char
            return int(estimated_chars * token_density)

        # === Excel / Word 文件 ===
        if file_extension in ["xlsx", "xls", "docx"]:
            return int(file_size / 6)

        # === 其他未知类型 ===
        return int(file_size / 8)

    except Exception as e:
        print(f"⚠️ Error estimating tokens: {e}")
        return 100

# ========================
# 实用函数
# ========================
def clear_token_cache():
    global _token_cache
    _token_cache.clear()

def get_cache_stats():
    return {"cache_size": len(_token_cache), "max_size": _cache_max_size}


if __name__ == '__main__':
    file_get = FileOperation()
    # content = file_get("./", ["9Q311103_(Token：54982).pdf"])
    # print(content)
